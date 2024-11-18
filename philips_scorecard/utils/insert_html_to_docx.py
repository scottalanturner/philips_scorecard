from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from bs4 import BeautifulSoup
import os

def set_cell_background(cell, hex_color):
    """Set background color of a cell"""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    )

def set_cell_border(cell, styles, border_color='4A5568'):
    """Set cell borders with specified color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create border elements
    for edge in ['top', 'left', 'bottom', 'right']:
        border = parse_xml(f'''<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                          <w:{edge} w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>
                        </w:tcBorders>''')
        tcPr.append(border)

def set_table_width(table, width_percent=100):
    """Set table width as percentage of page width"""
    table.table_direction = 1  # Set table to left-to-right
    table._tbl.xpath("./w:tblPr")[0].append(
        parse_xml(f'<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{width_percent*50}" w:type="pct"/>')
    )

def set_cell_padding(table, padding_twips=120):  # 120 twips = ~8.5 points
    """Set default cell padding for the table"""
    tblPr = table._element.xpath('./w:tblPr')
    if not tblPr:
        tblPr = parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        table._element.insert(0, tblPr)
    else:
        tblPr = tblPr[0]
        
    tblCellMar = parse_xml(f'''<w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:top w:w="{padding_twips}" w:type="dxa"/>
        <w:left w:w="{padding_twips}" w:type="dxa"/>
        <w:bottom w:w="{padding_twips}" w:type="dxa"/>
        <w:right w:w="{padding_twips}" w:type="dxa"/>
    </w:tblCellMar>''')
    tblPr.append(tblCellMar)

def convert_html_to_docx_elements(doc, html_content):
    """Convert HTML content to docx elements using BeautifulSoup"""
    soup = BeautifulSoup(html_content, 'html.parser')
    added_elements = []

    for element in soup.children:
        if element.name == 'table':
            # Create table
            rows = element.find_all('tr')
            if rows:
                cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                added_elements.append(table)

                # Set table to full width
                set_table_width(table)

                # Set default cell padding
                set_cell_padding(table)


                # Get table styles including border color
                table_style = element.get('style', '')
                if table_style:
                    table_styles = dict(s.strip().split(':') for s in table_style.split(';') 
                                    if ':' in s and s.strip())
                    border_color = None
                    if 'border' in table_styles:
                        # Extract color from border style like "1px solid #4A5568"
                        border_parts = table_styles['border'].split()
                        if len(border_parts) >= 3 and border_parts[2].startswith('#'):
                            border_color = border_parts[2].replace('#', '')

                # Get column widths from colgroup
                col_widths = []
                colgroup = element.find('colgroup')
                if colgroup:
                    for col in colgroup.find_all('col'):
                        width = col.get('width', '').rstrip('%')
                        col_widths.append(float(width) if width else None)

                # Process each row
                for i, row in enumerate(rows):
                    # Get row background color
                    bg_color = row.get('style', '')
                    bg_color = next((s.split(':')[1].strip() for s in bg_color.split(';') 
                                  if 'background-color' in s), None)

                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        # Get cell styles
                        cell_style = cell.get('style', '')
                        styles = dict(s.strip().split(':') for s in cell_style.split(';') 
                                   if ':' in s and s.strip())
                        
                        # Apply content and styling
                        table_cell = table.cell(i, (cols - 1) - j)
                        paragraph = table_cell.paragraphs[0]
                        
                        # Process cell contents for formatting
                        for content in cell.contents:
                            if isinstance(content, str):
                                # Preserve spaces in text, only strip from ends if needed
                                text = content.rstrip('\n')  # Remove newlines but keep spaces
                                if text:  # Only add if there's content (including spaces)
                                    run = paragraph.add_run(text)
                            elif content.name == 'b' or content.name == 'strong':
                                # Bold text
                                run = paragraph.add_run(content.get_text().strip())
                                run.bold = True
                            elif content.name == 'br':
                                # Add line break
                                paragraph.add_run().add_break()
                            elif content.name == 'ul':
                                # Handle unordered lists within cells
                                for li in content.find_all('li'):
                                    # Create a new paragraph for each list item
                                    list_para = table_cell.add_paragraph()
                                    # Add bullet character and text with proper spacing
                                    run = list_para.add_run('• ')
                                    run.font.symbol = True
                                    list_para.add_run(li.get_text().strip())
                                    # Add left indentation for list items
                                    list_para.paragraph_format.left_indent = Inches(0.25)
                                    list_para.paragraph_format.first_line_indent = Inches(-0.25)

                        # Apply borders only if specified
                        if any(border_prop in styles for border_prop in ['border', 'border-top', 'border-left', 'border-bottom', 'border-right']):
                            set_cell_border(table_cell, styles, border_color)
                        
                        # Apply background color - check cell first, then row
                        if 'background-color' in styles:
                            cell_bg = styles['background-color'].replace('#', '').strip()
                            set_cell_background(table_cell, cell_bg)
                        elif bg_color:  # fallback to row background if cell has none
                            bg_color = bg_color.replace('#', '').strip()
                            set_cell_background(table_cell, bg_color)
                            
                        # Apply text alignment
                        if 'text-align' in styles:
                            if styles['text-align'].strip() == 'center':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif styles['text-align'].strip() == 'right':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
                        # Apply font color if specified
                        if 'color' in styles:
                            color = styles['color'].replace('#', '').strip()
                            rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(*rgb)

                # Apply column widths
                if col_widths:
                    for i, width in enumerate(col_widths):
                        if width is not None:
                            for cell in table.columns[(cols - 1) - i].cells:
                                cell.width = Inches(width / 100 * 6)  # assuming 6 inches total width

        elif element.name == 'p':
            p = doc.add_paragraph()
            added_elements.append(p)
            
            # Handle paragraph styling
            p_style = element.get('style', '')
            if p_style:
                styles = dict(s.split(':') for s in p_style.split(';') 
                            if ':' in s and s.strip())
                
                # Apply text color
                if 'color' in styles:
                    color = styles['color'].replace('#', '').strip()
                    rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(*rgb)
                
                # Apply font style
                if 'font-style' in styles and styles['font-style'].strip() == 'italic':
                    for run in p.runs:
                        run.italic = True
            
            # Handle inline elements and their styling
            for child in element.children:
                if child.name in ['b', 'strong']:
                    run = p.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['i', 'em']:
                    run = p.add_run(child.get_text())
                    run.italic = True
                else:
                    text = child.string if child.string else child.get_text()
                    if text:
                        p.add_run(text)


    return added_elements


def update_doc_template_with_rtf(doc : Document, replacements) -> bool:
    """
    Replace placeholders in Word template with formatted HTML content
    
    Args:
        template_path (str): Path to Word template file
        replacements (dict): Dictionary of placeholder:html_content pairs
        output_path (str): Path to save the resulting document
    """
    try:
        
        # Process paragraphs where replacements are needed
        for paragraph in doc.paragraphs:
            for placeholder, html in replacements.items():
                if f"{{{{{placeholder}}}}}" in paragraph.text:
                    # Get the paragraph index
                    p_idx = doc._element.body.index(paragraph._element)
                    
                    # Convert HTML to docx elements
                    elements = convert_html_to_docx_elements(doc, html)
                    
                    # Insert elements at the correct position
                    for idx, element in enumerate(elements):
                        doc._element.body.insert(p_idx + idx, element._element)
                    
                    # Check if any element is a table by checking the XML tag
                    if any(element._element.tag.endswith('tbl') for element in elements):
                        empty_p = doc.add_paragraph()
                        doc._element.body.insert(p_idx + len(elements), empty_p._element)
                        
                    # Remove the placeholder paragraph
                    paragraph._element.getparent().remove(paragraph._element)
                    break  # Process only one replacement per paragraph
                
        return True
        
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        return False
    

# Example usage
if __name__ == "__main__":
    # Example with fully styled HTML input
    replacements = {
        "{{bp1}}": """
            <table style="width:100%; border-collapse:collapse; border:1px solid #4A5568;">
                <colgroup>
                    <col width="35%">
                    <col width="20%">
                    <col width="20%">
                    <col width="25%">
                </colgroup>
                <tr style="background-color:#2C5282;">
                    <th style="border:1px solid #4A5568; padding:8px; color:#FFFFFF; text-align:left;">Metric</th>
                    <th style="border:1px solid #4A5568; padding:8px; color:#FFFFFF; text-align:center;">Q1</th>
                    <th style="border:1px solid #4A5568; padding:8px; color:#FFFFFF; text-align:center;">Q2</th>
                    <th style="border:1px solid #4A5568; padding:8px; color:#FFFFFF; text-align:center;">Change</th>
                </tr>
                <tr style="background-color:#F7FAFC;">
                    <td style="border:1px solid #4A5568; padding:8px; font-weight:bold;">Revenue</td>
                    <td style="border:1px solid #4A5568; padding:8px; text-align:center;">$1.2M</td>
                    <td style="border:1px solid #4A5568; padding:8px; text-align:center;">$1.5M</td>
                    <td style="border:1px solid #4A5568; padding:8px; text-align:center; color:#48BB78;">+25%</td>
                </tr>
                <tr style="background-color:#EDF2F7;">
                    <td style="border:1px solid #4A5568; padding:8px; font-weight:bold;">Expenses</td>
                    <td style="border:1px solid #4A5568; padding:8px; text-align:center;">$800K</td>
                    <td style="border:1px solid #4A5568; padding:8px; text-align:center;">$850K</td>
                    <td style="border:1px solid #4A5568; padding:8px; text-align:center; color:#F56565;">+6.25%</td>
                </tr>
            </table>
            <p style="color:#718096; font-style:italic; margin-top:8px;">* All financial figures rounded to nearest 50K</p>
        """
    }
    doc = Document("template.docx")
    success = update_doc_template_with_rtf(
        doc=doc,
        replacements=replacements
    )
    if success:
        # Save the modified document
        doc.save('output.docx')
        print("Document saved successfully!")
    else:
        print("Error processing document.")


def replace_placeholders_in_docx(document : Document, replacements : str) -> str:

    success = update_doc_template_with_rtf(document, replacements)
    if not success:
        raise Exception("Error replacing placeholders in document")
